from __future__ import annotations

# ---------------------------------------------------------------------------
# ChromaDB / SQLite compatibility patch — MUST run before chromadb is
# imported anywhere. Streamlit Community Cloud's system sqlite3 is older
# than chromadb requires; pysqlite3-binary bundles a modern build.
# ---------------------------------------------------------------------------
import sys

try:
    __import__("pysqlite3")
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    # Running somewhere with a modern system sqlite3 already (e.g. some
    # local dev setups) — safe to continue without the shim.
    pass

import io
import json
import logging
import os
import random
import re
import time
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st

# `st.set_page_config` must be the first Streamlit command executed.
st.set_page_config(
    page_title="AI HR Recruitment Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============================================================================
# SECTION 1 — Pydantic models  (verbatim from models.py)
# ============================================================================

from pydantic import BaseModel, Field


class CandidateScreen(BaseModel):
    """Structured screening output for a single candidate."""

    candidate_name: str = Field(..., description="Candidate's full name or filename if unknown")
    file_name: str = Field(..., description="Original resume file name")
    matched_skills: List[str] = Field(default_factory=list, description="Skills that match the job description")
    missing_skills: List[str] = Field(default_factory=list, description="Required skills not found in resume")
    certifications: List[str] = Field(default_factory=list, description="Certifications held by the candidate")
    awards: List[str] = Field(default_factory=list, description="Awards or honors received")
    notable_projects: List[str] = Field(default_factory=list, description="Notable projects mentioned")
    experience_summary: str = Field(default="", description="Summary of relevant work experience")
    education_summary: str = Field(default="", description="Summary of education background")
    overall_impression: str = Field(default="", description="Screening agent's overall impression")


class ScreeningReport(BaseModel):
    candidates: List[CandidateScreen] = Field(default_factory=list)


class CandidateRanking(BaseModel):
    candidate_name: str
    file_name: str
    score: float = Field(..., ge=0, le=100, description="Overall candidate score out of 100")
    rank: int = Field(..., ge=1, description="Rank position, 1 = best")
    justification: str = Field(default="", description="Explanation for the assigned score/rank")
    confidence_level: str = Field(default="Medium", description="Low / Medium / High confidence in the ranking")
    fairness_note: str = Field(default="", description="Self-review note confirming the ranking was made without bias")


class RankingReport(BaseModel):
    rankings: List[CandidateRanking] = Field(default_factory=list)


class CandidateInterview(BaseModel):
    candidate_name: str
    file_name: str
    questions: List[str] = Field(default_factory=list, description="5-7 tailored interview questions")


class InterviewReport(BaseModel):
    interviews: List[CandidateInterview] = Field(default_factory=list)


class CandidateRecommendation(BaseModel):
    candidate_name: str
    file_name: str
    verdict: str = Field(..., description="Strong Hire / Hire / Maybe / No Hire")
    summary: str = Field(default="", description="Short summary supporting the verdict")
    key_strengths: List[str] = Field(default_factory=list)
    key_risks: List[str] = Field(default_factory=list)


class RecommendationReport(BaseModel):
    recommendations: List[CandidateRecommendation] = Field(default_factory=list)


class CandidateFullProfile(BaseModel):
    """Fully merged view of a single candidate across all four agents."""

    candidate_name: str
    file_name: str

    matched_skills: List[str] = Field(default_factory=list)
    missing_skills: List[str] = Field(default_factory=list)
    certifications: List[str] = Field(default_factory=list)
    awards: List[str] = Field(default_factory=list)
    notable_projects: List[str] = Field(default_factory=list)
    experience_summary: str = ""
    education_summary: str = ""
    overall_impression: str = ""

    score: Optional[float] = None
    rank: Optional[int] = None
    ranking_justification: str = ""
    confidence_level: str = ""
    fairness_note: str = ""

    interview_questions: List[str] = Field(default_factory=list)

    verdict: Optional[str] = None
    recommendation_summary: str = ""
    key_strengths: List[str] = Field(default_factory=list)
    key_risks: List[str] = Field(default_factory=list)


class AgentTraceStep(BaseModel):
    agent_name: str
    status: str = Field(default="completed", description="completed / failed / skipped")
    duration_seconds: Optional[float] = None
    message: str = ""


class FinalReport(BaseModel):
    job_description_excerpt: str = ""
    candidates_processed: int = 0
    candidates: List[CandidateFullProfile] = Field(default_factory=list)
    agent_trace: List[AgentTraceStep] = Field(default_factory=list)
    skipped_files: List[str] = Field(default_factory=list)
    errors: List[str] = Field(default_factory=list)


# ============================================================================
# SECTION 2 — Resume text extraction  (from main.py, unchanged logic)
# ============================================================================

from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_CANDIDATES = 5


def _extract_pdf_text(raw_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(raw_bytes))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages_text).strip()


def _extract_docx_text(raw_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(raw_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs).strip()


def _extract_txt_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="ignore").strip()


def extract_text(file_name: str, raw_bytes: bytes) -> str:
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".pdf":
        return _extract_pdf_text(raw_bytes)
    if ext == ".docx":
        return _extract_docx_text(raw_bytes)
    if ext == ".txt":
        return _extract_txt_text(raw_bytes)
    raise ValueError(f"Unsupported file type: {ext}")


# ============================================================================
# SECTION 3 — Local RAG tools: ChromaDB + LlamaIndex  (from tools.py)
#
# NOTE: the original tools.py kept the Chroma client / index / query engine
# in module-level globals. That's replaced here with st.session_state so
# each browser session gets its own isolated resume index — important once
# this runs on a shared Streamlit Cloud process instead of one FastAPI
# request per user.
# ============================================================================

import chromadb
from crewai.tools import tool
from llama_index.core import Document, Settings, StorageContext, VectorStoreIndex
from llama_index.core.query_engine import BaseQueryEngine
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore


@st.cache_resource(show_spinner="Loading local embedding model (first run only)...")
def get_embed_model() -> HuggingFaceEmbedding:
    """Lazily instantiate and cache the HuggingFace embedding model.

    Cached with st.cache_resource because the model itself is stateless
    and safe to share across all sessions in this process.
    """
    logger.info("Loading HuggingFace embedding model BAAI/bge-small-en-v1.5 ...")
    model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")
    Settings.embed_model = model
    # We only use LlamaIndex for retrieval, not generation.
    Settings.llm = None
    return model


def reset_index() -> None:
    """Reset this session's index state. Call before indexing a new batch."""
    st.session_state.chroma_client = None
    st.session_state.vector_index = None
    st.session_state.query_engine = None


def index_resumes(resumes: Dict[str, str]) -> VectorStoreIndex:
    """Build an in-memory ChromaDB-backed VectorStoreIndex from resume texts,
    scoped to the current Streamlit session.
    """
    embed_model = get_embed_model()

    client = chromadb.EphemeralClient()
    collection = client.get_or_create_collection("resumes")
    vector_store = ChromaVectorStore(chroma_collection=collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    documents: List[Document] = [
        Document(
            text=text,
            metadata={"file_name": file_name},
            excluded_llm_metadata_keys=["file_name"],
        )
        for file_name, text in resumes.items()
    ]

    index = VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        embed_model=embed_model,
    )

    st.session_state.chroma_client = client
    st.session_state.vector_index = index
    st.session_state.query_engine = None
    logger.info("Indexed %d resumes into ChromaDB.", len(documents))
    return index


def get_query_engine(similarity_top_k: int = 3):
    """Return a retriever over this session's currently built index."""
    index = st.session_state.get("vector_index")
    if index is None:
        raise RuntimeError("No resumes have been indexed yet. Call index_resumes() first.")

    if st.session_state.get("query_engine") is None:
        st.session_state.query_engine = index.as_retriever(similarity_top_k=similarity_top_k)
    return st.session_state.query_engine


def _format_nodes(nodes) -> str:
    if not nodes:
        return "No relevant resume content found."
    chunks = []
    for node_with_score in nodes:
        node = node_with_score.node
        file_name = node.metadata.get("file_name", "unknown")
        score = getattr(node_with_score, "score", None)
        score_str = f" (relevance: {score:.2f})" if score is not None else ""
        chunks.append(f"--- Resume: {file_name}{score_str} ---\n{node.get_content()}")
    return "\n\n".join(chunks)


@tool("Resume Retrieval Tool")
def resume_retrieval_tool(query: str) -> str:
    """
    Search the indexed resumes for content relevant to the given query.
    Use this tool to look up specific skills, experience, education,
    certifications, or projects mentioned across candidate resumes.

    Args:
        query: A natural language search query, e.g. "Python experience"
               or "AWS certifications".

    Returns:
        A formatted string containing the most relevant resume excerpts.
    """
    try:
        engine = get_query_engine()
        nodes = engine.retrieve(query)
        return _format_nodes(nodes)
    except RuntimeError as exc:
        return f"Resume retrieval unavailable: {exc}"
    except Exception as exc:  # noqa: BLE001
        logger.exception("Resume retrieval tool failed")
        return f"Resume retrieval failed due to an internal error: {exc}"


def run_resume_query(query: str, top_k: int = 3) -> List[Dict[str, Any]]:
    """In-process replacement for the old POST /query endpoint."""
    engine = get_query_engine(similarity_top_k=top_k)
    nodes = engine.retrieve(query)
    return [
        {
            "file_name": n.node.metadata.get("file_name", "unknown"),
            "score": getattr(n, "score", None),
            "text": n.node.get_content(),
        }
        for n in nodes
    ]


# ============================================================================
# SECTION 4 — CrewAI agents, Groq-backed  (from agents.py)
# ============================================================================

from crewai import LLM, Agent


def _patch_crewai_cache_breakpoint() -> None:
    """
    Monkey-patch crewai.llms.cache so that any dict/object it produces
    never carries a 'cache_breakpoint' key/attribute, which Groq's API
    does not accept.
    """
    try:
        import crewai.llms.cache as crewai_cache
    except ImportError:
        logger.warning("crewai.llms.cache module not found; skipping cache_breakpoint patch.")
        return

    for attr_name in dir(crewai_cache):
        attr = getattr(crewai_cache, attr_name, None)
        if not callable(attr):
            continue

        original_fn = attr

        def _make_wrapper(fn):
            def _wrapped(*args, **kwargs):
                result = fn(*args, **kwargs)
                if isinstance(result, dict) and "cache_breakpoint" in result:
                    result = {k: v for k, v in result.items() if k != "cache_breakpoint"}
                elif hasattr(result, "cache_breakpoint"):
                    try:
                        delattr(result, "cache_breakpoint")
                    except AttributeError:
                        pass
                return result
            return _wrapped

        try:
            setattr(crewai_cache, attr_name, _make_wrapper(original_fn))
        except (AttributeError, TypeError):
            continue

    logger.info("Applied CrewAI cache_breakpoint compatibility patch for Groq.")


_patch_crewai_cache_breakpoint()


@st.cache_resource(show_spinner=False)
def build_agents(groq_api_key: str) -> Dict[str, Agent]:
    """Construct and cache the 4 CrewAI agents for a given Groq API key.

    Cached per api_key value via st.cache_resource, so switching keys in
    the sidebar rebuilds agents but repeat runs with the same key reuse them.
    """
    llm = LLM(
        model="groq/llama-3.3-70b-versatile",
        api_key=groq_api_key,
        temperature=0.3,
        max_tokens=2000,
        max_retries=2,
    )

    screening_agent = Agent(
        role="Resume Screening Specialist",
        goal=(
            "Carefully analyze each candidate's resume against the job description "
            "and extract a structured, accurate profile covering matched and missing "
            "skills, certifications, awards, notable projects, experience, and education."
        ),
        backstory=(
            "You are a meticulous technical recruiter with over a decade of experience "
            "screening resumes for engineering and business roles. You read resumes "
            "closely, avoid assumptions not supported by the text, and always ground "
            "your findings in what is actually written in the candidate's resume."
        ),
        tools=[resume_retrieval_tool],
        llm=llm,
        allow_delegation=False,
        verbose=True,
    )

    ranking_agent = Agent(
        role="Candidate Ranking Analyst",
        goal=(
            "Objectively score and rank candidates from 0-100 based solely on how "
            "well their qualifications match the job requirements. Provide a clear "
            "justification for every score and perform an explicit fairness self-review "
            "to ensure the ranking is free of bias related to name, gender, age, "
            "ethnicity, or any factor unrelated to job qualifications."
        ),
        backstory=(
            "You are an experienced, impartial hiring analyst known for data-driven, "
            "defensible rankings. You explicitly double-check your own reasoning for "
            "signs of unconscious bias before finalizing any ranking, and you document "
            "that self-review as part of your output."
        ),
        tools=[],
        llm=llm,
        allow_delegation=False,
        verbose=True,
    )

    interview_agent = Agent(
        role="Interview Question Designer",
        goal=(
            "Design 5-7 tailored interview questions for each of the top 3 ranked "
            "candidates, mixing technical, behavioral, and situational questions "
            "that probe both the role's requirements and the candidate's specific "
            "background, including any gaps identified during screening."
        ),
        backstory=(
            "You are a senior interview panel lead who crafts precise, role-specific "
            "and candidate-specific questions that reveal real signal about a "
            "candidate's ability to succeed in the role."
        ),
        tools=[],
        llm=llm,
        allow_delegation=False,
        verbose=True,
    )

    recommendation_agent = Agent(
        role="Hiring Recommendation Lead",
        goal=(
            "Synthesize the screening and ranking results into a final hiring "
            "verdict for each candidate — one of 'Strong Hire', 'Hire', 'Maybe', "
            "or 'No Hire' — along with a concise summary, key strengths, and key risks."
        ),
        backstory=(
            "You are a hiring committee chair responsible for making the final "
            "call on every candidate. You weigh the evidence presented by the "
            "screening and ranking analysts and communicate clear, actionable "
            "verdicts to the hiring manager."
        ),
        tools=[],
        llm=llm,
        allow_delegation=False,
        verbose=True,
    )

    return {
        "screening": screening_agent,
        "ranking": ranking_agent,
        "interview": interview_agent,
        "recommendation": recommendation_agent,
    }


# ============================================================================
# SECTION 5 — Pipeline orchestration  (from tasks.py, adapted to take an
# `agents` dict instead of importing module-level agent singletons, and to
# accept an optional progress callback so the UI can update a progress bar
# between stages instead of polling an HTTP endpoint)
# ============================================================================

from crewai import Crew, Process, Task

INTER_AGENT_DELAY_RANGE = (5, 10)  # seconds
MAX_RETRIES = 3
JOB_DESC_MAX_CHARS = 500
ANALYZE_TIMEOUT_SECONDS = 300


def _truncate_job_description(job_description: str) -> str:
    text = job_description.strip()
    if len(text) <= JOB_DESC_MAX_CHARS:
        return text
    return text[:JOB_DESC_MAX_CHARS].rsplit(" ", 1)[0] + "..."


def _sleep_between_agents() -> None:
    delay = random.uniform(*INTER_AGENT_DELAY_RANGE)
    logger.info("Waiting %.1fs before next agent stage to respect rate limits...", delay)
    time.sleep(delay)


def _extract_json(raw_output: str):
    if raw_output is None:
        raise ValueError("Empty output from agent.")

    text = raw_output.strip()
    text = re.sub(r"^```(json)?", "", text.strip(), flags=re.IGNORECASE).strip()
    text = re.sub(r"```$", "", text.strip()).strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    for open_ch, close_ch in (("{", "}"), ("[", "]")):
        start = text.find(open_ch)
        end = text.rfind(close_ch)
        if start != -1 and end != -1 and end > start:
            candidate = text[start:end + 1]
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                continue

    raise ValueError(f"Could not parse JSON from agent output: {text[:300]}")


def _run_crew_with_retry(crew: Crew, stage_name: str) -> Tuple[Optional[str], Optional[str]]:
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            result = crew.kickoff()
            raw = str(result)
            return raw, None
        except Exception as exc:  # noqa: BLE001
            message = str(exc)
            is_rate_limit = "429" in message or "rate" in message.lower()
            if attempt < MAX_RETRIES and is_rate_limit:
                backoff = (2 ** attempt) + random.uniform(0, 2)
                logger.warning(
                    "[%s] Rate limit hit (attempt %d/%d). Backing off %.1fs.",
                    stage_name, attempt, MAX_RETRIES, backoff,
                )
                time.sleep(backoff)
                continue
            logger.exception("[%s] Agent stage failed on attempt %d.", stage_name, attempt)
            return None, message
    return None, "Max retries exceeded."


def run_screening_stage(
    agents: Dict[str, Agent], job_description: str, resumes: Dict[str, str]
) -> Tuple[List[CandidateScreen], AgentTraceStep]:
    started = time.time()
    jd_excerpt = _truncate_job_description(job_description)
    resume_list_text = "\n\n".join(f"File: {name}\n{text[:3000]}" for name, text in resumes.items())

    task = Task(
        description=(
            f"Job description (excerpt): {jd_excerpt}\n\n"
            f"Below are {len(resumes)} candidate resumes. For EACH resume, use the "
            "Resume Retrieval Tool if needed to confirm details, then extract a "
            "structured profile.\n\n"
            f"{resume_list_text}\n\n"
            "Return ONLY a JSON array, one object per candidate, with EXACTLY these "
            "keys: candidate_name, file_name, matched_skills (list), missing_skills "
            "(list), certifications (list), awards (list), notable_projects (list), "
            "experience_summary (string), education_summary (string), "
            "overall_impression (string). Do not include any text outside the JSON array."
        ),
        expected_output="A JSON array of candidate screening profiles.",
        agent=agents["screening"],
    )
    crew = Crew(agents=[agents["screening"]], tasks=[task], process=Process.sequential, verbose=True)

    raw, error = _run_crew_with_retry(crew, "Screening")
    duration = time.time() - started

    if error:
        return [], AgentTraceStep(agent_name="Resume Screening Agent", status="failed", duration_seconds=duration, message=error)

    try:
        data = _extract_json(raw)
        screens = [CandidateScreen(**item) for item in data]
        return screens, AgentTraceStep(
            agent_name="Resume Screening Agent", status="completed",
            duration_seconds=duration, message=f"Screened {len(screens)} candidate(s).",
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to parse screening output.")
        return [], AgentTraceStep(agent_name="Resume Screening Agent", status="failed", duration_seconds=duration, message=f"Parse error: {exc}")


def run_ranking_stage(
    agents: Dict[str, Agent], job_description: str, screens: List[CandidateScreen]
) -> Tuple[List[CandidateRanking], AgentTraceStep]:
    started = time.time()
    jd_excerpt = _truncate_job_description(job_description)
    profiles_json = json.dumps([s.model_dump() for s in screens])

    task = Task(
        description=(
            f"Job description (excerpt): {jd_excerpt}\n\n"
            f"Candidate screening profiles:\n{profiles_json}\n\n"
            "Score each candidate 0-100 based only on job fit, then rank them "
            "(1 = best). For each candidate, write a brief justification, a "
            "confidence_level ('Low', 'Medium', or 'High'), and a fairness_note "
            "explicitly confirming your score was based only on job-relevant "
            "qualifications and not on name, gender, age, or ethnicity.\n\n"
            "Return ONLY a JSON array with EXACTLY these keys per object: "
            "candidate_name, file_name, score (number), rank (integer), "
            "justification (string), confidence_level (string), fairness_note (string)."
        ),
        expected_output="A JSON array of candidate rankings sorted best to worst.",
        agent=agents["ranking"],
    )
    crew = Crew(agents=[agents["ranking"]], tasks=[task], process=Process.sequential, verbose=True)

    raw, error = _run_crew_with_retry(crew, "Ranking")
    duration = time.time() - started

    if error:
        return [], AgentTraceStep(agent_name="Candidate Ranking Agent", status="failed", duration_seconds=duration, message=error)

    try:
        data = _extract_json(raw)
        rankings = [CandidateRanking(**item) for item in data]
        rankings.sort(key=lambda r: r.rank)
        return rankings, AgentTraceStep(
            agent_name="Candidate Ranking Agent", status="completed",
            duration_seconds=duration, message=f"Ranked {len(rankings)} candidate(s).",
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to parse ranking output.")
        return [], AgentTraceStep(agent_name="Candidate Ranking Agent", status="failed", duration_seconds=duration, message=f"Parse error: {exc}")


def run_interview_stage(
    agents: Dict[str, Agent],
    job_description: str,
    rankings: List[CandidateRanking],
    screens_by_file: Dict[str, CandidateScreen],
) -> Tuple[List[CandidateInterview], AgentTraceStep]:
    started = time.time()
    jd_excerpt = _truncate_job_description(job_description)

    top_candidates = rankings[:3]
    if not top_candidates:
        return [], AgentTraceStep(agent_name="Interview Question Generator", status="skipped", duration_seconds=0.0, message="No ranked candidates available.")

    context_items = []
    for r in top_candidates:
        screen = screens_by_file.get(r.file_name)
        context_items.append({
            "candidate_name": r.candidate_name,
            "file_name": r.file_name,
            "matched_skills": screen.matched_skills if screen else [],
            "missing_skills": screen.missing_skills if screen else [],
            "experience_summary": screen.experience_summary if screen else "",
        })

    task = Task(
        description=(
            f"Job description (excerpt): {jd_excerpt}\n\n"
            f"Top candidates:\n{json.dumps(context_items)}\n\n"
            "For EACH top candidate, write 5 to 7 interview questions mixing "
            "technical, behavioral, and situational types. Questions must be "
            "specific to both the role and the candidate's background "
            "(e.g. probe missing skills, validate matched skills).\n\n"
            "Return ONLY a JSON array with EXACTLY these keys per object: "
            "candidate_name, file_name, questions (list of strings)."
        ),
        expected_output="A JSON array of interview question sets for the top candidates.",
        agent=agents["interview"],
    )
    crew = Crew(agents=[agents["interview"]], tasks=[task], process=Process.sequential, verbose=True)

    raw, error = _run_crew_with_retry(crew, "Interview")
    duration = time.time() - started

    if error:
        return [], AgentTraceStep(agent_name="Interview Question Generator", status="failed", duration_seconds=duration, message=error)

    try:
        data = _extract_json(raw)
        interviews = [CandidateInterview(**item) for item in data]
        return interviews, AgentTraceStep(
            agent_name="Interview Question Generator", status="completed",
            duration_seconds=duration, message=f"Generated questions for {len(interviews)} top candidate(s).",
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to parse interview output.")
        return [], AgentTraceStep(agent_name="Interview Question Generator", status="failed", duration_seconds=duration, message=f"Parse error: {exc}")


def run_recommendation_stage(
    agents: Dict[str, Agent], job_description: str, rankings: List[CandidateRanking]
) -> Tuple[List[CandidateRecommendation], AgentTraceStep]:
    started = time.time()
    jd_excerpt = _truncate_job_description(job_description)
    rankings_json = json.dumps([r.model_dump() for r in rankings])

    task = Task(
        description=(
            f"Job description (excerpt): {jd_excerpt}\n\n"
            f"Candidate rankings:\n{rankings_json}\n\n"
            "For EACH candidate, issue a final hiring verdict: one of "
            "'Strong Hire', 'Hire', 'Maybe', or 'No Hire'. Provide a concise "
            "summary, key_strengths (list), and key_risks (list).\n\n"
            "Return ONLY a JSON array with EXACTLY these keys per object: "
            "candidate_name, file_name, verdict, summary, key_strengths (list), "
            "key_risks (list)."
        ),
        expected_output="A JSON array of hiring recommendations for all candidates.",
        agent=agents["recommendation"],
    )
    crew = Crew(agents=[agents["recommendation"]], tasks=[task], process=Process.sequential, verbose=True)

    raw, error = _run_crew_with_retry(crew, "Recommendation")
    duration = time.time() - started

    if error:
        return [], AgentTraceStep(agent_name="Hiring Recommendation Agent", status="failed", duration_seconds=duration, message=error)

    try:
        data = _extract_json(raw)
        recs = [CandidateRecommendation(**item) for item in data]
        return recs, AgentTraceStep(
            agent_name="Hiring Recommendation Agent", status="completed",
            duration_seconds=duration, message=f"Produced recommendations for {len(recs)} candidate(s).",
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Failed to parse recommendation output.")
        return [], AgentTraceStep(agent_name="Hiring Recommendation Agent", status="failed", duration_seconds=duration, message=f"Parse error: {exc}")


def run_pipeline(
    agents: Dict[str, Agent],
    job_description: str,
    resumes: Dict[str, str],
    skipped_files: Optional[List[str]] = None,
    progress_callback: Optional[Any] = None,
) -> FinalReport:
    """Run the full sequential 4-agent pipeline and merge results into a
    single FinalReport. `progress_callback(stage_index, stage_label)` is
    called after each stage completes, for driving a Streamlit progress bar.
    """
    trace: List[AgentTraceStep] = []
    errors: List[str] = []

    def _report_progress(idx: int, label: str) -> None:
        if progress_callback:
            progress_callback(idx, label)

    # --- Stage 1: Screening ---
    screens, step1 = run_screening_stage(agents, job_description, resumes)
    trace.append(step1)
    if step1.status == "failed":
        errors.append(f"Screening stage failed: {step1.message}")
    screens_by_file = {s.file_name: s for s in screens}
    _report_progress(1, "Screening complete")

    if screens:
        _sleep_between_agents()

    # --- Stage 2: Ranking ---
    rankings, step2 = run_ranking_stage(agents, job_description, screens) if screens else (
        [], AgentTraceStep(agent_name="Candidate Ranking Agent", status="skipped", message="No screened candidates to rank.")
    )
    trace.append(step2)
    if step2.status == "failed":
        errors.append(f"Ranking stage failed: {step2.message}")
    _report_progress(2, "Ranking complete")

    if rankings:
        _sleep_between_agents()

    # --- Stage 3: Interview Questions (top 3) ---
    interviews, step3 = run_interview_stage(agents, job_description, rankings, screens_by_file) if rankings else (
        [], AgentTraceStep(agent_name="Interview Question Generator", status="skipped", message="No rankings available.")
    )
    trace.append(step3)
    if step3.status == "failed":
        errors.append(f"Interview stage failed: {step3.message}")
    interviews_by_file = {i.file_name: i for i in interviews}
    _report_progress(3, "Interview questions drafted")

    if rankings:
        _sleep_between_agents()

    # --- Stage 4: Hiring Recommendation ---
    recommendations, step4 = run_recommendation_stage(agents, job_description, rankings) if rankings else (
        [], AgentTraceStep(agent_name="Hiring Recommendation Agent", status="skipped", message="No rankings available.")
    )
    trace.append(step4)
    if step4.status == "failed":
        errors.append(f"Recommendation stage failed: {step4.message}")
    recs_by_file = {r.file_name: r for r in recommendations}
    _report_progress(4, "Recommendations finalized")

    # --- Merge everything programmatically ---
    rankings_by_file = {r.file_name: r for r in rankings}
    all_file_names = set(screens_by_file) | set(rankings_by_file) | set(recs_by_file)

    profiles: List[CandidateFullProfile] = []
    for file_name in all_file_names:
        screen = screens_by_file.get(file_name)
        ranking = rankings_by_file.get(file_name)
        interview = interviews_by_file.get(file_name)
        rec = recs_by_file.get(file_name)

        name = (
            (screen.candidate_name if screen else None)
            or (ranking.candidate_name if ranking else None)
            or (rec.candidate_name if rec else None)
            or file_name
        )

        profiles.append(CandidateFullProfile(
            candidate_name=name,
            file_name=file_name,
            matched_skills=screen.matched_skills if screen else [],
            missing_skills=screen.missing_skills if screen else [],
            certifications=screen.certifications if screen else [],
            awards=screen.awards if screen else [],
            notable_projects=screen.notable_projects if screen else [],
            experience_summary=screen.experience_summary if screen else "",
            education_summary=screen.education_summary if screen else "",
            overall_impression=screen.overall_impression if screen else "",
            score=ranking.score if ranking else None,
            rank=ranking.rank if ranking else None,
            ranking_justification=ranking.justification if ranking else "",
            confidence_level=ranking.confidence_level if ranking else "",
            fairness_note=ranking.fairness_note if ranking else "",
            interview_questions=interview.questions if interview else [],
            verdict=rec.verdict if rec else None,
            recommendation_summary=rec.summary if rec else "",
            key_strengths=rec.key_strengths if rec else [],
            key_risks=rec.key_risks if rec else [],
        ))

    profiles.sort(key=lambda p: (p.rank is None, p.rank if p.rank is not None else 0))

    return FinalReport(
        job_description_excerpt=_truncate_job_description(job_description),
        candidates_processed=len(resumes),
        candidates=profiles,
        agent_trace=trace,
        skipped_files=skipped_files or [],
        errors=errors,
    )


# ============================================================================
# SECTION 6 — Streamlit frontend  (from app.py, adapted to call run_pipeline
# directly instead of POSTing to a FastAPI backend)
# ============================================================================

from reportlab.lib import colors as rl_colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

PRIMARY = "#6C2BD9"
PRIMARY_DARK = "#5522B0"
SUCCESS = "#10B981"
WARNING = "#F59E0B"
DANGER = "#EF4444"

TEXT_PRIMARY = "#0F172A"
TEXT_SECONDARY = "#475569"
TEXT_MUTED = "#94A3B8"

VERDICT_STYLE = {
    "Strong Hire": {"color": "#FFFFFF", "bg": "#10B981", "emoji": "🌟"},
    "Hire":        {"color": "#FFFFFF", "bg": "#059669", "emoji": "✅"},
    "Maybe":       {"color": "#FFFFFF", "bg": "#F59E0B", "emoji": "🤔"},
    "No Hire":     {"color": "#FFFFFF", "bg": "#EF4444", "emoji": "⛔"},
}

RANK_EMOJI = {1: "🥇", 2: "🥈", 3: "🥉"}

TECH_STACK = [
    {"icon": "⚡", "name": "Groq · Llama 3.3 70B", "detail": "High-performance LLM inference"},
    {"icon": "🤖", "name": "CrewAI", "detail": "Multi-agent orchestration"},
    {"icon": "🔍", "name": "LlamaIndex + ChromaDB", "detail": "RAG & vector search"},
    {"icon": "🎨", "name": "Streamlit", "detail": "Single-file app — no separate backend"},
]

AGENT_PIPELINE = [
    {"num": "01", "icon": "🔍", "name": "Screening Agent", "detail": "Extracts skills, experience, education"},
    {"num": "02", "icon": "📊", "name": "Ranking Agent", "detail": "Scores & ranks candidates (0-100)"},
    {"num": "03", "icon": "🎤", "name": "Interview Agent", "detail": "Generates questions for top 3"},
    {"num": "04", "icon": "✅", "name": "Recommendation Agent", "detail": "Final hire / no-hire verdict"},
]


def html_block(s: str) -> str:
    """Dedent multi-line HTML so Streamlit's Markdown parser doesn't mistake
    indented lines for a fenced code block (CommonMark rule)."""
    return "\n".join(line.strip() for line in s.strip("\n").splitlines())


st.markdown(
    f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
    html, body, [class*="css"] {{ font-family: 'Inter', sans-serif; }}
    .block-container {{ padding-top: 4rem; max-width: 1080px; }}
    section[data-testid="stSidebar"] {{ background-color: #FFFFFF; border-right: 1px solid #E2E8F0; box-shadow: 2px 0 8px rgba(15, 23, 42, 0.03); }}
    .sidebar-header {{ background: #FFFFFF; border: 1.5px solid #CBD5E1; border-radius: 14px; padding: 0.9rem 1rem; margin-bottom: 1.1rem; box-shadow: 0 2px 8px rgba(15, 23, 42, 0.05); }}
    .sidebar-header .title {{ font-size: 1.25rem; font-weight: 800; color: {TEXT_PRIMARY}; }}
    .sidebar-header .subtitle {{ font-size: 0.78rem; color: {TEXT_MUTED}; margin-top: 0.15rem; }}
    .sidebar-section-title {{ font-size: 0.8rem; font-weight: 800; text-transform: uppercase; letter-spacing: 0.05em; color: #000000; margin: 0.4rem 0 0.6rem 0; }}
    .tech-card {{ background: #FFFFFF; border: 1.5px solid #CBD5E1; border-radius: 12px; padding: 0.65rem 0.8rem; margin-bottom: 0.55rem; transition: transform 0.15s ease, box-shadow 0.15s ease; }}
    .tech-card:hover {{ transform: translateY(-2px); box-shadow: 0 6px 16px rgba(108, 43, 217, 0.12); border-color: {PRIMARY}; }}
    .tech-card .t-name {{ font-weight: 700; font-size: 0.85rem; color: {TEXT_PRIMARY}; }}
    .tech-card .t-detail {{ font-size: 0.73rem; color: {TEXT_SECONDARY}; margin-top: 0.1rem; }}
    .agent-card {{ background: #FAF7FF; border: 1.5px solid #D8BEF5; border-radius: 12px; padding: 0.65rem 0.8rem; margin-bottom: 0.55rem; display: flex; gap: 0.6rem; align-items: flex-start; }}
    .agent-card .a-num {{ font-weight: 800; font-size: 0.78rem; color: {PRIMARY}; background: #EFE4FC; border-radius: 6px; padding: 0.1rem 0.4rem; min-width: 26px; text-align: center; }}
    .agent-card .a-name {{ font-weight: 700; font-size: 0.85rem; color: {TEXT_PRIMARY}; }}
    .agent-card .a-detail {{ font-size: 0.73rem; color: {TEXT_SECONDARY}; margin-top: 0.1rem; }}
    .hero {{ background: linear-gradient(135deg, #F3EEFC 0%, #EDE4FA 100%); border: 1.5px solid #D8BEF5; border-radius: 32px; padding: 2.1rem 2.3rem; margin-bottom: 1.2rem; width: 100%; box-sizing: border-box; overflow: hidden; }}
    .hero h1 {{ font-size: 2rem; font-weight: 800; color: {TEXT_PRIMARY}; margin-bottom: 0.5rem; }}
    .hero .tagline {{ font-size: 1rem; color: {TEXT_SECONDARY}; margin: 0; }}
    .tech-box {{ display: inline-flex; align-items: center; gap: 0.4rem; background: linear-gradient(120deg, {PRIMARY}, {PRIMARY_DARK}); border: 1.5px solid {PRIMARY_DARK}; border-radius: 999px; padding: 0.35rem 0.85rem; margin: 0.2rem 0.3rem 0.2rem 0; font-size: 0.82rem; font-weight: 600; color: #FFFFFF; }}
    .section-label {{ font-size: 0.95rem; font-weight: 700; color: {TEXT_PRIMARY}; margin: 1.1rem 0 0.5rem 0; }}
    .card {{ background: #FFFFFF; border: 1.5px solid #CBD5E1; border-radius: 16px; padding: 1.3rem 1.5rem; box-shadow: 0 2px 10px rgba(15, 23, 42, 0.04); margin-bottom: 1.1rem; }}
    .metric-card {{ background: #FFFFFF; border: 1.5px solid #CBD5E1; border-radius: 16px; padding: 1.2rem 1rem; text-align: center; box-shadow: 0 2px 10px rgba(15, 23, 42, 0.04); }}
    .metric-card .m-icon {{ font-size: 1.5rem; }}
    .metric-card .m-value {{ font-size: 1.7rem; font-weight: 800; color: {PRIMARY}; margin-top: 0.2rem; }}
    .metric-card .m-label {{ font-size: 0.78rem; color: {TEXT_MUTED}; text-transform: uppercase; letter-spacing: 0.04em; margin-top: 0.15rem; }}
    .cmp-table {{ width: 100%; border-collapse: collapse; background: #FFFFFF; border: 1.5px solid #CBD5E1; border-radius: 14px; overflow: hidden; box-shadow: 0 2px 10px rgba(15, 23, 42, 0.05); }}
    .cmp-table th {{ background: {PRIMARY}; color: #FFFFFF; text-align: left; padding: 0.7rem 0.9rem; font-size: 0.82rem; text-transform: uppercase; letter-spacing: 0.03em; }}
    .cmp-table td {{ padding: 0.7rem 0.9rem; font-size: 0.9rem; color: {TEXT_PRIMARY}; border-bottom: 1px solid #EEF2F7; }}
    .cmp-table tr:nth-child(even) {{ background: #FAFBFD; }}
    .cmp-table tr:hover {{ background: #F3EEFC; }}
    .cmp-table a.cand-link {{ color: {TEXT_PRIMARY}; font-weight: 700; text-decoration: none; }}
    .cmp-table a.cand-link:hover {{ color: {PRIMARY}; text-decoration: underline; }}
    .pill {{ display: inline-block; padding: 0.22rem 0.7rem; border-radius: 999px; font-weight: 700; font-size: 0.78rem; }}
    .cand-card {{ background: #FFFFFF; border: 1.5px solid #CBD5E1; border-radius: 18px; padding: 1.6rem 1.7rem; margin-bottom: 1.3rem; box-shadow: 0 3px 14px rgba(15, 23, 42, 0.05); }}
    .cand-header {{ display: flex; align-items: center; justify-content: space-between; border-bottom: 1px solid #EEF2F7; padding-bottom: 0.9rem; margin-bottom: 0.9rem; }}
    .cand-name {{ font-size: 1.35rem; font-weight: 800; color: {TEXT_PRIMARY}; }}
    .cand-file {{ font-size: 0.8rem; color: {TEXT_MUTED}; margin-top: 0.15rem; }}
    .ring {{ width: 68px; height: 68px; border-radius: 50%; display: flex; align-items: center; justify-content: center; }}
    .ring-inner {{ width: 54px; height: 54px; border-radius: 50%; background: #FFFFFF; display: flex; align-items: center; justify-content: center; font-weight: 800; font-size: 0.95rem; }}
    .field-label {{ font-weight: 700; font-size: 0.88rem; color: {TEXT_PRIMARY}; margin: 0.85rem 0 0.35rem 0; }}
    .field-text {{ font-size: 0.88rem; color: {TEXT_SECONDARY}; line-height: 1.5; }}
    .field-block {{ padding-bottom: 0.9rem; margin-bottom: 0.9rem; border-bottom: 1px solid #E2E8F0; }}
    .field-block:last-child {{ border-bottom: none; margin-bottom: 0; padding-bottom: 0; }}
    .field-block.interview {{ margin-top: 0.4rem; }}
    .tag {{ display: inline-block; padding: 0.2rem 0.6rem; margin: 0.15rem 0.3rem 0.15rem 0; border-radius: 8px; font-size: 0.78rem; font-weight: 600; }}
    .tag.match {{ background: #D1FAE5; color: #047857; }}
    .tag.missing {{ background: #FEE2E2; color: #B91C1C; }}
    .tag.cert {{ background: #E0E7FF; color: #4338CA; }}
    .tag.strength {{ background: #D1FAE5; color: #047857; }}
    .tag.risk {{ background: #FEF3C7; color: #92400E; }}
    div.stButton > button {{ border-radius: 12px; font-weight: 700; padding: 0.65rem 1.2rem; border: none; }}
    div.stButton > button[kind="primary"] {{ background: linear-gradient(120deg, {PRIMARY}, {PRIMARY_DARK}); color: white; }}
    hr {{ margin: 1.4rem 0; border-color: #E2E8F0; }}
    </style>
    """,
    unsafe_allow_html=True,
)

# --- Session state ---
if "report" not in st.session_state:
    st.session_state.report = None
if "vector_index" not in st.session_state:
    st.session_state.vector_index = None
if "chroma_client" not in st.session_state:
    st.session_state.chroma_client = None
if "query_engine" not in st.session_state:
    st.session_state.query_engine = None


def get_groq_api_key() -> str:
    """Resolve the Groq API key: Streamlit secrets first (recommended for
    a deployed app), falling back to a sidebar input for local testing.
    """
    try:
        secret_key = st.secrets.get("GROQ_API_KEY", "")
    except Exception:  # noqa: BLE001 — no secrets.toml present locally
        secret_key = ""

    if secret_key:
        return secret_key

    return st.session_state.get("manual_groq_key", "")


# --- Sidebar ---
with st.sidebar:
    st.markdown(
        '<div class="sidebar-header"><div class="title">🤖 AI HR Assistant</div>'
        '<div class="subtitle">AI-Powered Recruitment Intelligence</div></div>',
        unsafe_allow_html=True,
    )

    if not get_groq_api_key():
        st.markdown('<div class="sidebar-section-title">Groq API Key</div>', unsafe_allow_html=True)
        st.session_state.manual_groq_key = st.text_input(
            "Groq API key", type="password", label_visibility="collapsed",
            placeholder="gsk_...", help="Get a free key at console.groq.com. "
            "On Streamlit Cloud, set this as a GROQ_API_KEY secret instead.",
        )
        st.caption("💡 On Streamlit Cloud, add `GROQ_API_KEY` under App settings → Secrets instead of pasting it here.")

    st.markdown('<div class="sidebar-section-title">Technology Stack</div>', unsafe_allow_html=True)
    for tech in TECH_STACK:
        st.markdown(
            f'<div class="tech-card"><div class="t-name">{tech["icon"]} {tech["name"]}</div>'
            f'<div class="t-detail">{tech["detail"]}</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown('<div class="sidebar-section-title">Agent Pipeline</div>', unsafe_allow_html=True)
    for agent in AGENT_PIPELINE:
        st.markdown(
            f'<div class="agent-card"><div class="a-num">{agent["num"]}</div><div>'
            f'<div class="a-name">{agent["icon"]} {agent["name"]}</div>'
            f'<div class="a-detail">{agent["detail"]}</div></div></div>',
            unsafe_allow_html=True,
        )

# --- Hero ---
st.markdown(
    '<div class="hero"><h1>🤖 AI HR Recruitment Assistant</h1>'
    '<p class="tagline">A multi-agent AI pipeline that screens, ranks, and evaluates candidate resumes end to end.</p></div>',
    unsafe_allow_html=True,
)
st.markdown(
    "".join(f'<span class="tech-box">{t["icon"]} {t["name"]}</span>' for t in TECH_STACK),
    unsafe_allow_html=True,
)
st.write("")

# --- Input section ---
st.markdown('<div class="section-label">📝 Job Description</div>', unsafe_allow_html=True)
job_description = st.text_area(
    "Job description", height=180, label_visibility="collapsed",
    placeholder="Paste the job description here — e.g. We're hiring a Senior Backend "
    "Engineer with 5+ years of experience in Python, distributed systems, and AWS...",
)

st.markdown('<div class="section-label">📁 Upload Resumes (up to 5)</div>', unsafe_allow_html=True)
max_slider = st.slider("Number of candidates to process", min_value=1, max_value=MAX_CANDIDATES, value=MAX_CANDIDATES)
uploaded_files = st.file_uploader("Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True, label_visibility="collapsed")
if uploaded_files and len(uploaded_files) > max_slider:
    st.warning(f"Only the first {max_slider} file(s) will be analyzed based on the slider above.")
    uploaded_files = uploaded_files[:max_slider]
if uploaded_files:
    for f in uploaded_files:
        st.caption(f"✅ {f.name} · {f.size / 1024:.1f} KB")

st.write("")
groq_key_available = bool(get_groq_api_key())
run_clicked = st.button(
    "🚀 Run Analysis", type="primary", use_container_width=True,
    disabled=not (job_description and uploaded_files and groq_key_available),
)
if not groq_key_available:
    st.caption("⚠️ Add a Groq API key in the sidebar (or as a secret) to enable analysis.")
elif not job_description or not uploaded_files:
    st.caption("Add a job description and at least one resume to enable analysis.")
st.markdown("---")


# --- In-process pipeline runner (replaces the old POST /analyze call) ---

def run_analysis_in_process(job_desc: str, files: List[Any], groq_api_key: str) -> Optional[Dict[str, Any]]:
    resume_texts: Dict[str, str] = {}
    skipped_files: List[str] = []

    for f in files:
        ext = os.path.splitext(f.name or "")[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            skipped_files.append(f.name or "unknown")
            continue
        try:
            raw_bytes = f.getvalue()
            text = extract_text(f.name, raw_bytes)
            if not text.strip():
                skipped_files.append(f.name)
                continue
            resume_texts[f.name] = text
        except Exception:  # noqa: BLE001
            logger.exception("Failed to parse file %s", f.name)
            skipped_files.append(f.name or "unknown")

    if not resume_texts:
        st.error("No valid resume text could be extracted from the uploaded files.")
        return None

    progress = st.progress(0, text="Screening resumes...")
    stage_labels = {1: "Ranking candidates...", 2: "Drafting interview questions...", 3: "Finalizing recommendations...", 4: "Done"}

    def _on_progress(stage_idx: int, _label: str) -> None:
        pct = min(int(stage_idx / 4 * 100), 100)
        progress.progress(pct, text=stage_labels.get(stage_idx, "Working..."))

    start = time.time()
    try:
        reset_index()
        index_resumes(resume_texts)
        agents = build_agents(groq_api_key)

        report = run_pipeline(
            agents, job_desc, resume_texts, skipped_files,
            progress_callback=_on_progress,
        )
    except Exception as exc:  # noqa: BLE001
        progress.empty()
        logger.exception("Pipeline execution failed.")
        st.error(f"Pipeline execution failed: {exc}")
        return None
    finally:
        progress.empty()

    result = report.model_dump()
    result["_client_elapsed_seconds"] = time.time() - start
    return result


if run_clicked:
    with st.spinner("Running the 4-agent pipeline — this can take a couple of minutes..."):
        result = run_analysis_in_process(job_description, uploaded_files, get_groq_api_key())
    if result:
        st.session_state.report = result
        st.toast("Analysis complete!", icon="✅")


# --- Derived metric helpers ---

def skill_match_pct(c: Dict[str, Any]) -> float:
    matched, missing = len(c.get("matched_skills") or []), len(c.get("missing_skills") or [])
    total = matched + missing
    return round((matched / total) * 100) if total else 0.0


def skill_match_color(pct: float) -> str:
    return SUCCESS if pct > 70 else WARNING if pct >= 40 else DANGER


def rank_label(rank: Optional[int]) -> str:
    return "—" if rank is None else RANK_EMOJI.get(rank, f"#{rank}")


def processing_time_str(report: Dict[str, Any]) -> str:
    durations = [s.get("duration_seconds") for s in report.get("agent_trace", []) if s.get("duration_seconds")]
    total = sum(durations) if durations else report.get("_client_elapsed_seconds", 0)
    if not total:
        return "—"
    return f"{total:.0f}s" if total < 60 else f"{total / 60:.1f}m"


# --- Result renderers ---

def render_dashboard(report: Dict[str, Any], candidates: List[Dict[str, Any]]) -> None:
    scores = [c["score"] for c in candidates if c.get("score") is not None]
    avg_score = f"{(sum(scores) / len(scores)):.0f}" if scores else "—"
    strong_hires = sum(1 for c in candidates if c.get("verdict") == "Strong Hire")

    metrics = [
        ("👥", len(candidates), "Total Candidates"),
        ("🏆", strong_hires, "Strong Hires"),
        ("📊", avg_score, "Average Score"),
        ("⚡", processing_time_str(report), "Processing Time"),
    ]
    cols = st.columns(4)
    for col, (icon, value, label) in zip(cols, metrics):
        col.markdown(
            f'<div class="metric-card"><div class="m-icon">{icon}</div>'
            f'<div class="m-value">{value}</div><div class="m-label">{label}</div></div>',
            unsafe_allow_html=True,
        )


def render_comparison_table(candidates: List[Dict[str, Any]]) -> None:
    rows = []
    for idx, c in enumerate(candidates):
        rank = c.get("rank")
        score = c.get("score")
        verdict = c.get("verdict") or "Pending"
        vstyle = VERDICT_STYLE.get(verdict, {"color": TEXT_PRIMARY, "bg": "#E2E8F0", "emoji": "•"})
        pct = skill_match_pct(c)
        pct_color = skill_match_color(pct)
        score_txt = f"{score:.0f}/100" if score is not None else "—"

        rows.append(
            f'<tr><td>{rank_label(rank)}</td>'
            f'<td><a class="cand-link" href="#cand-{idx}">{c["candidate_name"]}</a></td>'
            f'<td>{score_txt}</td>'
            f'<td><span class="pill" style="background:{vstyle["bg"]}; color:{vstyle["color"]};">'
            f'{vstyle["emoji"]} {verdict}</span></td>'
            f'<td><span style="color:{pct_color}; font-weight:700;">{pct}%</span></td></tr>'
        )

    table_html = (
        '<table class="cmp-table"><thead><tr><th>Rank</th><th>Candidate</th><th>Score</th>'
        f'<th>Verdict</th><th>Skill Match %</th></tr></thead><tbody>{"".join(rows)}</tbody></table>'
    )
    st.markdown(table_html, unsafe_allow_html=True)


def _tags(items: List[str], css_class: str, empty_text: str = "None recorded") -> str:
    if not items:
        return f'<span class="field-text">{empty_text}</span>'
    return "".join(f'<span class="tag {css_class}">{i}</span>' for i in items)


def _field_block(icon: str, label: str, content: str, extra_class: str = "") -> str:
    cls = f"field-block {extra_class}".strip()
    return f'<div class="{cls}"><div class="field-label">{icon} {label}</div><div class="field-text">{content}</div></div>'


def _candidate_fields_html(c: Dict[str, Any]) -> str:
    fields = [
        ("📝", "Recommendation Summary", c.get("recommendation_summary") or "No recommendation summary available."),
        ("✅", "Strengths", _tags(c.get("key_strengths", []), "strength")),
        ("⚠️", "Risks", _tags(c.get("key_risks", []), "risk")),
        ("🛠️", "Matched Skills", _tags(c.get("matched_skills", []), "match")),
        ("❌", "Missing Skills", _tags(c.get("missing_skills", []), "missing")),
        ("🎓", "Certifications", _tags(c.get("certifications", []), "cert")),
        ("🏗️", "Notable Projects", _tags(c.get("notable_projects", []), "cert")),
        ("💼", "Experience", c.get("experience_summary") or "Not available."),
        ("🎓", "Education", c.get("education_summary") or "Not available."),
        ("💭", "Overall Impression", c.get("overall_impression") or "Not available."),
        ("⚖️", "Ranking Justification", c.get("ranking_justification") or "Not available."),
    ]
    html = "".join(_field_block(icon, label, content) for icon, label, content in fields)
    questions = "".join(f"{i}. {q}<br/><br/>" for i, q in enumerate(c.get("interview_questions", []), 1)) \
        or "Not generated for this candidate."
    html += _field_block("🎤", "Interview Questions", questions, "interview")
    return html


def render_candidate_profiles(candidates: List[Dict[str, Any]]) -> None:
    for idx, c in enumerate(candidates):
        rank, score = c.get("rank"), c.get("score")
        verdict = c.get("verdict") or "Pending"
        vstyle = VERDICT_STYLE.get(verdict, {"color": TEXT_PRIMARY, "bg": "#E2E8F0", "emoji": "•"})
        ring_color = skill_match_color(score) if score is not None else TEXT_MUTED
        score_display = f"{score:.0f}" if score is not None else "—"
        deg = (score or 0) * 3.6

        st.markdown(
            html_block(
                f"""
                <div class="cand-card" id="cand-{idx}">
                    <div class="cand-header">
                        <div>
                            <div class="cand-name">{rank_label(rank)} {c['candidate_name']}</div>
                            <div class="cand-file">{c['file_name']}</div>
                            <span class="pill" style="background:{vstyle['bg']}; color:{vstyle['color']}; margin-top:0.5rem; display:inline-block;">
                                {vstyle['emoji']} {verdict}</span>
                        </div>
                        <div class="ring" style="background: conic-gradient({ring_color} {deg}deg, #E5E7EB 0deg);">
                            <div class="ring-inner" style="color:{ring_color};">{score_display}</div>
                        </div>
                    </div>
                    {_candidate_fields_html(c)}
                </div>
                """
            ),
            unsafe_allow_html=True,
        )


# --- PDF export ---

def generate_pdf(report: Dict[str, Any], candidates: List[Dict[str, Any]]) -> bytes:
    def pdf_rank(rank: Optional[int]) -> str:
        return f"#{rank}" if rank is not None else "—"

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=1.8 * cm, rightMargin=1.8 * cm, topMargin=1.6 * cm, bottomMargin=1.6 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], textColor=rl_colors.HexColor(PRIMARY))
    h2_style = ParagraphStyle("H2X", parent=styles["Heading2"], textColor=rl_colors.HexColor(TEXT_PRIMARY), spaceBefore=10)
    h3_style = ParagraphStyle("H3X", parent=styles["Heading3"], textColor=rl_colors.HexColor(PRIMARY), spaceBefore=8)
    body_style = ParagraphStyle("BodyX", parent=styles["BodyText"], textColor=rl_colors.HexColor(TEXT_SECONDARY), leading=14)

    story = [
        Paragraph("AI HR Recruitment Assistant — Report", title_style),
        Paragraph("The right candidate isn't just found — they're understood.", body_style),
        Spacer(1, 0.5 * cm),
    ]
    if report.get("job_description_excerpt"):
        story += [Paragraph("Job Description (excerpt)", h2_style), Paragraph(report["job_description_excerpt"], body_style), Spacer(1, 0.4 * cm)]

    story.append(Paragraph("Candidate Comparison", h2_style))
    table_data = [["Rank", "Candidate", "Score", "Verdict", "Skill Match %"]]
    for c in candidates:
        score_txt = f"{c['score']:.0f}/100" if c.get("score") is not None else "—"
        table_data.append([pdf_rank(c.get("rank")), c.get("candidate_name", ""), score_txt, c.get("verdict") or "Pending", f"{skill_match_pct(c)}%"])
    tbl = Table(table_data, hAlign="LEFT", colWidths=[2 * cm, 5.5 * cm, 2.5 * cm, 3 * cm, 3 * cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor(PRIMARY)), ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#E2E8F0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#FAFBFD")]), ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.6 * cm))

    sections_map = [
        ("Recommendation Summary", "recommendation_summary", False), ("Strengths", "key_strengths", True),
        ("Risks", "key_risks", True), ("Matched Skills", "matched_skills", True), ("Missing Skills", "missing_skills", True),
        ("Certifications", "certifications", True), ("Notable Projects", "notable_projects", True),
        ("Experience", "experience_summary", False), ("Education", "education_summary", False),
        ("Overall Impression", "overall_impression", False), ("Ranking Justification", "ranking_justification", False),
    ]
    for c in candidates:
        block = [Paragraph(f"{pdf_rank(c.get('rank'))} {c.get('candidate_name', '')}", h2_style)]
        score_line = (
            f"File: {c.get('file_name', '')} | Score: {c['score']:.0f}/100"
            if c.get("score") is not None else f"File: {c.get('file_name', '')}"
        )
        block.append(Paragraph(score_line, body_style))
        block.append(Paragraph(f"Verdict: {c.get('verdict') or 'Pending'}", body_style))

        for label, key, is_list in sections_map:
            value = ", ".join(c.get(key, [])) if is_list else c.get(key)
            if value:
                block.append(Paragraph(label, h3_style))
                block.append(Paragraph(value, body_style))

        if c.get("interview_questions"):
            block.append(Paragraph("Interview Questions", h3_style))
            for i, q in enumerate(c["interview_questions"], 1):
                block.append(Paragraph(f"{i}. {q}", body_style))
        block.append(Spacer(1, 0.5 * cm))
        story.append(KeepTogether(block))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# --- Results ---

report = st.session_state.report

if report is None:
    st.markdown(
        '<div class="card" style="text-align:center; padding:2.6rem 1.5rem;">'
        '<div style="font-size:2.2rem;">📥</div>'
        '<div style="font-weight:700; font-size:1.1rem; margin-top:0.4rem;">No analysis yet</div>'
        '<div style="color:#94A3B8; margin-top:0.3rem;">Fill in the job description and upload resumes '
        'above, then click <strong>Run Analysis</strong>.</div></div>',
        unsafe_allow_html=True,
    )
else:
    candidates = sorted(report.get("candidates", []), key=lambda c: (c.get("rank") is None, c.get("rank") if c.get("rank") is not None else 0))

    st.markdown('<div class="section-label">📊 Dashboard</div>', unsafe_allow_html=True)
    render_dashboard(report, candidates)
    if report.get("skipped_files"):
        st.warning("Skipped files: " + ", ".join(report["skipped_files"]))
    for err in report.get("errors", []):
        st.error(err)

    st.markdown('<div class="section-label">📋 Candidate Comparison</div>', unsafe_allow_html=True)
    if candidates:
        render_comparison_table(candidates)
    else:
        st.info("No candidates to compare.")

    st.markdown('<div class="section-label">👤 Candidate Profiles</div>', unsafe_allow_html=True)
    if candidates:
        render_candidate_profiles(candidates)
    else:
        st.info("No candidate profiles available.")

    st.markdown("---")
    pdf_bytes = generate_pdf(report, candidates)
    st.download_button("⬇️ Download PDF Report", data=pdf_bytes, file_name="hr_recruitment_report.pdf", mime="application/pdf", use_container_width=True)
